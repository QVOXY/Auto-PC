import os
import shutil
from docx import Document

# 构建路径
path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'para_library', 'transmit2'))

try:
    # 读取文件内容
    with open(path, 'r') as f:
        num_list = [int(line.strip()) for line in f.readlines()]
except FileNotFoundError:
    print(f"文件 {path} 未找到。")
    exit(1)
except ValueError:
    print(f"文件 {path} 中的内容不是有效的整数。")
    exit(1)

# 获取最后一个数字
if num_list:
    num = num_list[-1]
else:
    print("文件中没有数字。")
    exit(1)

# 构建源文件路径
original_file_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'procedure', 'process1', f'paragraph_{num+1}.docx'))

# 检查源文件是否存在
if not os.path.exists(original_file_path):
    print(f"源文件 {original_file_path} 不存在。")
    exit(1)

# 构建目标文件夹路径
destination_folder_path = os.path.abspath(os.path.join(os.path.dirname(__file__), '..', 'procedure', 'process2'))

# 新的文件名
new_file_name = 'References-.docx'

# 确保目标文件夹存在
if not os.path.exists(destination_folder_path):
    os.makedirs(destination_folder_path)

# 目标文件的完整路径
destination_file_path = os.path.join(destination_folder_path, new_file_name)

try:
    # 打开源文件
    doc = Document(original_file_path)
    
    # 在文档头部添加 "References-"
    if doc.paragraphs:
        doc.paragraphs[0].insert_paragraph_before('References-')
    else:
        doc.add_paragraph('References-')
    
    # 保存修改后的文档到目标路径
    doc.save(destination_file_path)
    
    print(f"成功生成文件: {new_file_name}")
except shutil.Error as e:
    print(f"复制文件时出错: {e}")
except Exception as e:
    print(f"处理文件时出错: {e}")
